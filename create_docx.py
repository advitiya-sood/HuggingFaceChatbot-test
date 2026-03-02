import os
from docx import Document
import re

def markdown_to_docx(md_filepath, docx_filepath):
    if not os.path.exists(md_filepath):
        print(f"Error: {md_filepath} not found.")
        return

    with open(md_filepath, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    doc.add_heading('Chatbot FAQ Responses', 0)

    # Simple parsing to convert markdown elements to docx paragraphs
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('#'):
            # Skip the main header as we've added it manually
            if not line.startswith('# Chatbot FAQ'):
                level = min(len(line) - len(line.lstrip('#')), 9)
                text = line.lstrip('#').strip()
                doc.add_heading(text, level=level if level <= 9 else 9)
        elif line == '---':
            doc.add_paragraph('--------------------------------------------------')
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '')).bold = True
        else:
            # Handle basic bold tags within a line
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    try:
        doc.save(docx_filepath)
        print(f"Successfully created {docx_filepath}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == '__main__':
    md_file = "faq_responses.md"
    docx_file = "faq_responses.docx"
    markdown_to_docx(md_file, docx_file)
